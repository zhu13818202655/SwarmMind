"""Docx renderer built directly with ``python-docx``.

We do not use docxtpl in M1: keeping templates as pure Python style dicts
(see :mod:`docx_styles`) avoids checking binary ``.docx`` files into git and
gives us full control over preset variants. The renderer assembles the
document programmatically from :class:`ReportContext`.
"""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from swarmmind.domains.fly_report.export.base import (
    BaseRenderer,
    RenderedArtifact,
)
from swarmmind.domains.fly_report.export.docx_styles import (
    DocxStyle,
    get_style,
)
from swarmmind.domains.fly_report.export.template_loader import LoadedTemplate
from swarmmind.domains.fly_report.schemas import (
    OutputFormat,
    ReportBlock,
    ReportContext,
    ReportSection,
)

DOCX_FONT_NAME = "FangSong"
DOCX_BODY_SIZE_PT = 11


class DocxRenderer(BaseRenderer):
    output_format: OutputFormat = "docx"

    def render_markdown(
        self,
        markdown: str,
        *,
        output_dir: Path,
        filename: str = "markdown-report.docx",
        template_ref: str | None = None,
        title: str | None = None,
    ) -> RenderedArtifact:
        loaded = self._template_loader.load(
            output_format=self.output_format,
            template_ref=template_ref,
        )
        output_dir.mkdir(parents=True, exist_ok=True)

        out_path = output_dir / _normalize_docx_filename(filename)
        _render_markdown_with_pandoc(
            markdown,
            output_path=out_path,
            title=title,
        )
        return RenderedArtifact(
            output_format=self.output_format,
            artifact_path=str(out_path),
            template_ref=f"markdown:{loaded.template_ref}",
        )

    def _render(
        self,
        ctx: ReportContext,
        *,
        loaded: LoadedTemplate,
        output_dir: Path,
    ) -> RenderedArtifact:
        style = get_style(loaded.name)
        doc = Document()
        _apply_page_setup(doc, style)
        _write_title(doc, ctx, style)
        _write_meta(doc, ctx, style)
        _write_sections(doc, ctx, style, charts_dir=output_dir / "charts")
        _write_footer(doc, ctx, style)
        _apply_docx_font_policy(doc)

        out_path = output_dir / f"{ctx.session_id}.docx"
        doc.save(str(out_path))
        return RenderedArtifact(
            output_format=self.output_format,
            artifact_path=str(out_path),
            template_ref=loaded.template_ref,
        )


# ---------------------------------------------------------------------- helpers


def _render_markdown_with_pandoc(
    markdown: str,
    *,
    output_path: Path,
    title: str | None,
) -> None:
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        raise RuntimeError("Pandoc is required for Markdown-to-DOCX rendering but was not found")

    normalized = _normalize_markdown_for_pandoc(markdown, title=title)
    source_path = output_path.with_suffix(".pandoc.md")
    source_path.write_text(normalized, encoding="utf-8")

    command = [
        pandoc,
        str(source_path),
        "--from",
        "markdown+fenced_divs+bracketed_spans+link_attributes+pipe_tables+raw_html",
        "--to",
        "docx",
        "--output",
        str(output_path),
    ]
    result = subprocess.run(
        command,
        cwd=str(source_path.parent),
        capture_output=True,
        text=True,
        check=False,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(
            "Pandoc Markdown-to-DOCX rendering failed: "
            f"{(result.stderr or result.stdout).strip()}"
        )
    _postprocess_pandoc_docx(output_path, normalized)
    source_path.unlink(missing_ok=True)


def _normalize_markdown_for_pandoc(markdown: str, *, title: str | None) -> str:
    lines: list[str] = []
    if title:
        lines.extend([f"# {title}", ""])
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if align := _alignment_directive(line):
            lines.append(f"::: {{align={align}}}")
            continue
        lines.append(raw_line)
    return "\n".join(lines).strip() + "\n"


def _postprocess_pandoc_docx(path: Path, markdown: str) -> None:
    doc = Document(str(path))
    _apply_pandoc_alignment(doc, markdown)
    _apply_pandoc_table_layouts(doc, markdown)
    _apply_pandoc_image_alignment(doc, markdown)
    _apply_pandoc_colors(doc, markdown)
    _apply_docx_font_policy(doc)
    doc.save(str(path))


def _apply_pandoc_alignment(doc: Document, markdown: str) -> None:
    paragraph_alignments = _pandoc_paragraph_alignments(markdown)
    if not paragraph_alignments:
        return
    paragraph_index = 0
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip() and not paragraph._p.xpath(".//w:drawing"):
            continue
        if paragraph_index < len(paragraph_alignments):
            _apply_paragraph_alignment(paragraph, paragraph_alignments[paragraph_index])
        paragraph_index += 1


def _pandoc_paragraph_alignments(markdown: str) -> list[str]:
    alignments: list[str] = []
    current_align = "left"
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if align := _pandoc_alignment_directive(line):
            current_align = align
            continue
        if line == ":::":
            current_align = "left"
            continue
        if not line or _is_markdown_table_separator(line):
            continue
        if _is_markdown_table_row(line):
            continue
        alignments.append(current_align)
    return alignments


def _pandoc_alignment_directive(line: str) -> str | None:
    match = re.fullmatch(r":::\s*\{align=(left|center|right|justify)\}", line)
    return match.group(1) if match else _alignment_directive(line)


def _apply_pandoc_colors(doc: Document, markdown: str) -> None:
    for text, color in _colored_text_runs(markdown):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if text and text in run.text:
                    run.font.color.rgb = RGBColor(*color)


def _colored_text_runs(markdown: str) -> list[tuple[str, tuple[int, int, int]]]:
    pairs: list[tuple[str, tuple[int, int, int]]] = []
    for match in re.finditer(r"\[(?P<text>[^\]]+)\]\{color=#(?P<hex>[0-9a-fA-F]{6})\}", markdown):
        color = _parse_rgb(match.group("hex"))
        if color:
            pairs.append((match.group("text"), color))
    for match in re.finditer(
        r"<span\s+style=[\"']color:\s*#(?P<hex>[0-9a-fA-F]{6});?[\"']>(?P<text>.*?)</span>",
        markdown,
    ):
        color = _parse_rgb(match.group("hex"))
        if color:
            pairs.append((match.group("text"), color))
    return pairs


def _apply_pandoc_table_layouts(doc: Document, markdown: str) -> None:
    alignments, pairs = _pandoc_table_layouts(markdown)
    tables = list(doc.tables)
    paired_indices = {index for pair in pairs for index in pair}
    for index, align in enumerate(alignments):
        if index >= len(tables):
            break
        tables[index].alignment = _table_alignment(align)
        _style_pandoc_table(tables[index], paired=index in paired_indices)
        if index not in paired_indices:
            _set_table_width_percent(tables[index], _adaptive_table_width_percent(tables[index]))

    for left_index, right_index in reversed(pairs):
        if left_index >= len(tables) or right_index >= len(tables):
            continue
        _wrap_tables_side_by_side(doc, tables[left_index], tables[right_index])

    _add_spacing_after_top_level_tables(doc)


def _pandoc_table_layouts(markdown: str) -> tuple[list[str], list[tuple[int, int]]]:
    alignments: list[str] = []
    pairs: list[tuple[int, int]] = []
    stack: list[dict[str, Any]] = []
    current_align = "left"
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if attrs := _pandoc_div_attrs(line):
            entry: dict[str, Any] = {
                "previous_align": current_align,
                "is_table_pair": _is_table_pair_attrs(attrs),
                "tables": [],
            }
            if attrs.get("align") in {"left", "center", "right", "justify"}:
                current_align = attrs["align"]
            stack.append(entry)
            index += 1
            continue
        if line == ":::" and stack:
            entry = stack.pop()
            if entry["is_table_pair"]:
                table_indices = entry["tables"]
                for pair_index in range(0, len(table_indices) - 1, 2):
                    pairs.append((table_indices[pair_index], table_indices[pair_index + 1]))
            current_align = entry["previous_align"]
            index += 1
            continue
        if _is_markdown_table_start(lines, index):
            table_index = len(alignments)
            alignments.append(current_align)
            for entry in stack:
                if entry["is_table_pair"]:
                    entry["tables"].append(table_index)
            index += 1
            while index < len(lines) and _is_markdown_table_row(lines[index].strip()):
                index += 1
            continue
        index += 1
    return alignments, pairs


def _pandoc_div_attrs(line: str) -> dict[str, str] | None:
    if align := _alignment_directive(line):
        return {"align": align}
    match = re.fullmatch(r":::\s*\{(?P<attrs>[^}]*)\}", line)
    if not match:
        return None
    attrs = _parse_attrs(match.group("attrs"))
    classes = [part[1:] for part in match.group("attrs").split() if part.startswith(".")]
    if classes:
        attrs["class"] = " ".join(classes)
    return attrs


def _is_table_pair_attrs(attrs: dict[str, str]) -> bool:
    classes = set((attrs.get("class") or "").split())
    return (
        "table-pair" in classes
        or attrs.get("layout") in {"two-tables", "two-column-tables"}
        or attrs.get("table-layout") in {"pair", "two-column", "two-columns"}
    )


def _table_alignment(align: str):
    return {
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }.get(align, WD_TABLE_ALIGNMENT.LEFT)


def _adaptive_table_width_percent(table) -> int:
    column_count = len(table.columns)
    if column_count <= 2:
        return 72
    if column_count == 3:
        return 86
    return 96


def _style_pandoc_table(table, *, paired: bool = False) -> None:
    table.autofit = True
    _set_table_layout(table, "autofit")
    _set_table_cell_margins(table, top=90, bottom=90, left=120, right=120)
    _set_table_borders(table, color="C9D3DF")

    if not table.rows:
        return

    header = table.rows[0]
    header.height = None
    for cell in header.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                _format_table_run(run, bold=True, color=(31, 78, 121))

    for row_index, row in enumerate(table.rows[1:], start=1):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 0 and not paired:
                _shade_cell(cell, "F7FAFC")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _format_table_run(run, bold=False, color=(34, 34, 34))


def _format_table_run(run, *, bold: bool, color: tuple[int, int, int]) -> None:
    run.font.bold = bold
    _set_run_font(run, DOCX_FONT_NAME, DOCX_BODY_SIZE_PT)
    run.font.color.rgb = RGBColor(*color)


def _set_table_layout(table, layout_type: str) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), layout_type)


def _set_table_cell_margins(
    table,
    *,
    top: int,
    bottom: int,
    left: int,
    right: int,
) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for edge, value in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_borders(table, *, color: str) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _wrap_tables_side_by_side(doc: Document, left_table, right_table) -> None:
    parent = doc.add_table(rows=1, cols=2)
    parent.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width_percent(parent, 100)
    _set_table_borders_none(parent)
    for cell in parent.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_width_percent(cell, 50)
        cell._tc.clear_content()

    parent_element = parent._tbl
    parent_element.getparent().remove(parent_element)
    left_table._tbl.addprevious(parent_element)
    left_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    right_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_width_percent(left_table, 92)
    _set_table_width_percent(right_table, 92)
    parent.rows[0].cells[0]._tc.append(left_table._tbl)
    parent.rows[0].cells[1]._tc.append(right_table._tbl)


def _set_table_width_percent(table, percent: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(int(percent * 50)))


def _set_cell_width_percent(cell, percent: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "pct")
    tc_w.set(qn("w:w"), str(int(percent * 50)))


def _add_spacing_after_top_level_tables(doc: Document) -> None:
    for table in reversed(doc.tables):
        next_element = table._tbl.getnext()
        if next_element is not None and next_element.tag == qn("w:p"):
            _set_paragraph_spacing(next_element, after_twips=160)
            continue
        paragraph = OxmlElement("w:p")
        _set_paragraph_spacing(paragraph, after_twips=160)
        table._tbl.addnext(paragraph)


def _set_paragraph_spacing(paragraph_element, *, after_twips: int) -> None:
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:after"), str(after_twips))


def _set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _apply_pandoc_image_alignment(doc: Document, markdown: str) -> None:
    image_specs = _pandoc_image_specs(markdown)
    if not image_specs:
        return

    spec_index = 0
    paragraphs = list(doc.paragraphs)
    for paragraph_index, paragraph in enumerate(paragraphs):
        if not paragraph._p.xpath(".//w:drawing"):
            continue
        if spec_index >= len(image_specs):
            break
        caption, align = image_specs[spec_index]
        _apply_paragraph_alignment(paragraph, align)

        next_index = paragraph_index + 1
        if next_index < len(paragraphs):
            next_paragraph = paragraphs[next_index]
            if caption and next_paragraph.text.strip() == caption:
                _apply_paragraph_alignment(next_paragraph, align)
        spec_index += 1


def _pandoc_image_specs(markdown: str) -> list[tuple[str, str]]:
    specs: list[tuple[str, str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if image := _image_match(line):
            attrs = _parse_attrs(image.group("attrs") or "")
            specs.append((image.group("caption") or "", attrs.get("align", "center")))
    return specs


def _apply_page_setup(doc: Document, style: DocxStyle) -> None:
    for section in doc.sections:
        section.top_margin = Cm(style.page_margin_cm)
        section.bottom_margin = Cm(style.page_margin_cm)
        section.left_margin = Cm(style.page_margin_cm)
        section.right_margin = Cm(style.page_margin_cm)


def _write_title(doc: Document, ctx: ReportContext, style: DocxStyle) -> None:
    p = doc.add_paragraph()
    run = p.add_run(_title_text(ctx))
    _apply_run(run, style.title_font, style.title_size_pt, style.title_color_rgb, bold=True)


def _write_meta(doc: Document, ctx: ReportContext, style: DocxStyle) -> None:
    f = ctx.filter
    parts = [
        f"周期：{f.period.label}",
        f"范围：{_scope_label(f.dimension.scope)}",
    ]
    if f.dimension.department_ids:
        parts.append("部门：" + ", ".join(f.dimension.department_ids))
    if f.dimension.pilot_ids:
        parts.append("飞手：" + ", ".join(f.dimension.pilot_ids))
    parts.append(f"修订 v{ctx.revision}")
    p = doc.add_paragraph()
    run = p.add_run(" · ".join(parts))
    _apply_run(run, style.body_font, style.body_size_pt, (102, 102, 102))


def _write_sections(
    doc: Document,
    ctx: ReportContext,
    style: DocxStyle,
    *,
    charts_dir: Path,
) -> None:
    for index, section in enumerate(ctx.sections, start=1):
        _write_section(doc, section, style, charts_dir=charts_dir, index=index)


def _write_section(
    doc: Document,
    section: ReportSection,
    style: DocxStyle,
    *,
    charts_dir: Path,
    index: int | None = None,
) -> None:
    heading = doc.add_paragraph()
    heading_text = (
        f"{style.section_prefix}{section.title}"
        if style.section_prefix or index is None
        else f"{index}. {section.title}"
    )
    run = heading.add_run(heading_text)
    _apply_run(run, style.title_font, style.heading_size_pt, style.heading_color_rgb, bold=True)

    if section.blocks:
        for block in section.blocks:
            _write_block(doc, block, style, charts_dir=charts_dir)
    else:
        _write_legacy_section_body(doc, section, style, charts_dir=charts_dir)

    for child_index, child in enumerate(section.children, start=1):
        _write_section(doc, child, style, charts_dir=charts_dir, index=child_index)


def _write_legacy_section_body(
    doc: Document,
    section: ReportSection,
    style: DocxStyle,
    *,
    charts_dir: Path,
) -> None:
    if section.summary_md:
        p = doc.add_paragraph()
        run = p.add_run(section.summary_md)
        _apply_run(run, style.body_font, style.body_size_pt, (51, 51, 51))

    if section.kpis:
        _write_kpi_table(doc, section.kpis, style)

    for table in section.tables:
        _write_data_table(doc, table, style, caption=str(table.get("title") or "数据表"))

    for chart in section.charts:
        _write_chart(doc, chart, style, charts_dir=charts_dir, caption=chart.title)


def _write_block(
    doc: Document,
    block: ReportBlock,
    style: DocxStyle,
    *,
    charts_dir: Path,
) -> None:
    if block.style.page_break_before:
        doc.add_page_break()

    if block.kind == "heading":
        doc.add_heading(block.text, level=block.level)
    elif block.kind == "paragraph":
        p = doc.add_paragraph()
        _apply_paragraph_alignment(p, block.style.align)
        if block.runs:
            for text_run in block.runs:
                run = p.add_run(text_run.text)
                _apply_run(
                    run,
                    style.body_font,
                    style.body_size_pt,
                    _parse_rgb(text_run.color) or (51, 51, 51),
                    bold=text_run.bold,
                )
                run.italic = text_run.italic
                run.underline = text_run.underline
        else:
            run = p.add_run(block.text or "")
            _apply_run(run, style.body_font, style.body_size_pt, (51, 51, 51))
    elif block.kind == "markdown":
        _write_markdown(doc, block.markdown, style)
    elif block.kind == "list":
        for item in block.items:
            _write_list_item(doc, item, style, ordered=block.ordered, level=0)
    elif block.kind == "table":
        _write_data_table(doc, block.table, style, caption=block.caption or block.title)
    elif block.kind == "chart":
        _write_chart(doc, block.chart, style, charts_dir=charts_dir, caption=block.caption)
    elif block.kind == "chart_text":
        _write_chart(doc, block.chart, style, charts_dir=charts_dir, caption=block.caption)
        if block.text:
            p = doc.add_paragraph()
            run = p.add_run(block.text)
            _apply_run(run, style.body_font, style.body_size_pt, (51, 51, 51))
    elif block.kind == "image":
        image_path = Path(block.uri)
        if image_path.exists():
            width = Inches(block.width / 96) if block.width else Inches(6.0)
            doc.add_picture(str(image_path), width=width)
            if block.caption:
                _write_caption(doc, block.caption, style)
    elif block.kind == "kpi_group":
        _write_kpi_table(doc, block.kpis, style)
    elif block.kind == "callout":
        p = doc.add_paragraph()
        run = p.add_run(block.markdown)
        _apply_run(run, style.body_font, style.body_size_pt, (51, 51, 51), bold=True)
    elif block.kind == "page_break":
        doc.add_page_break()
    elif block.kind == "spacer":
        for _ in range(block.height):
            doc.add_paragraph()


def _write_markdown(doc: Document, markdown: str, style: DocxStyle) -> None:
    _write_markdown_document(doc, markdown, style)


def _write_markdown_document(doc: Document, markdown: str, style: DocxStyle) -> None:
    lines = markdown.splitlines()
    index = 0
    current_align = "left"
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue
        if align := _alignment_directive(line):
            current_align = align
            index += 1
            continue
        if line == ":::":
            current_align = "left"
            index += 1
            continue
        if _is_markdown_table_start(lines, index):
            index = _write_markdown_table(doc, lines, index, style)
            continue
        if image := _image_match(line):
            _write_markdown_image(doc, image, style, base_dir=None)
            index += 1
            continue
        if line in {"---", "***", "___"}:
            p = doc.add_paragraph("─" * 24)
            _apply_paragraph_alignment(p, current_align)
            index += 1
            continue
        if heading := _heading_match(line):
            paragraph = doc.add_heading(heading.group("text"), level=min(len(heading.group("marks")), 4))
            _apply_paragraph_alignment(paragraph, current_align)
        elif line.startswith(('- ', '* ')):
            p = doc.add_paragraph(style="List Bullet")
            _apply_paragraph_alignment(p, current_align)
            _write_inline_runs(p, line[2:], style)
        elif match := _ordered_list_match(line):
            p = doc.add_paragraph(style="List Number")
            _apply_paragraph_alignment(p, current_align)
            _write_inline_runs(p, match.group(1), style)
        else:
            p = doc.add_paragraph()
            _apply_paragraph_alignment(p, current_align)
            _write_inline_runs(p, line, style)
        index += 1


def _write_inline_runs(paragraph, text: str, style: DocxStyle) -> None:
    pattern = re.compile(
        r"(\[[^\]]+\]\{color=#[0-9a-fA-F]{6}\}|"
        r"<span\s+style=[\"']color:\s*#[0-9a-fA-F]{6};?[\"']>.*?</span>|"
        r"\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)"
    )
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            _apply_run(run, style.body_font, style.body_size_pt, (51, 51, 51))
        token = match.group(0)
        color = _inline_color(token)
        clean = _clean_inline_token(token)
        run = paragraph.add_run(clean)
        _apply_run(
            run,
            style.body_font,
            style.body_size_pt,
            color or (51, 51, 51),
            bold=token.startswith(("**", "__")),
        )
        run.italic = token.startswith(("*", "_")) and not token.startswith(("**", "__"))
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _apply_run(run, style.body_font, style.body_size_pt, (51, 51, 51))


def _clean_inline_token(token: str) -> str:
    color_link = re.fullmatch(r"\[(?P<text>[^\]]+)\]\{color=#[0-9a-fA-F]{6}\}", token)
    if color_link:
        return color_link.group("text")
    color_span = re.fullmatch(
        r"<span\s+style=[\"']color:\s*#[0-9a-fA-F]{6};?[\"']>(?P<text>.*?)</span>",
        token,
    )
    if color_span:
        return color_span.group("text")
    return token.strip("*_`")


def _inline_color(token: str) -> tuple[int, int, int] | None:
    match = re.search(r"#(?P<hex>[0-9a-fA-F]{6})", token)
    if not match or not (token.startswith("[") or token.startswith("<span")):
        return None
    return _parse_rgb(match.group("hex"))


def _alignment_directive(line: str) -> str | None:
    match = re.fullmatch(r":::\s*align=(left|center|right|justify)", line)
    return match.group(1) if match else None


def _heading_match(line: str):
    return re.match(r"^(?P<marks>#{1,6})\s+(?P<text>.+)$", line)


def _image_match(line: str):
    return re.fullmatch(
        r"!\[(?P<caption>[^\]]*)\]\((?P<path>[^)]+)\)(?:\{(?P<attrs>[^}]*)\})?",
        line,
    )


def _write_markdown_image(doc: Document, image, style: DocxStyle, *, base_dir: Path | None) -> None:
    attrs = _parse_attrs(image.group("attrs") or "")
    image_path = Path(image.group("path")).expanduser()
    if base_dir is not None and not image_path.is_absolute():
        image_path = base_dir / image_path
    if not image_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[missing image: {image_path}]")
        _apply_run(run, style.body_font, style.body_size_pt, (180, 40, 40), bold=True)
        return

    p = doc.add_paragraph()
    _apply_paragraph_alignment(p, attrs.get("align", "center"))
    run = p.add_run()
    width = _image_width(attrs)
    if width is None:
        run.add_picture(str(image_path), width=Inches(5.5))
    else:
        run.add_picture(str(image_path), width=width)
    caption = image.group("caption") or attrs.get("title")
    if caption:
        _write_caption(doc, caption, style)


def _parse_attrs(raw: str) -> dict[str, str]:
    attrs: dict[str, str] = {}
    for part in raw.split():
        if "=" not in part:
            continue
        key, value = part.split("=", 1)
        attrs[key.strip().lower()] = value.strip().strip('"\'')
    return attrs


def _image_width(attrs: dict[str, str]):
    raw = attrs.get("width") or attrs.get("w")
    if not raw:
        return None
    match = re.fullmatch(r"(?P<value>\d+(?:\.\d+)?)(?P<unit>cm|in|px)?", raw.strip())
    if not match:
        return None
    value = float(match.group("value"))
    unit = match.group("unit") or "px"
    if unit == "cm":
        return Cm(value)
    if unit == "in":
        return Inches(value)
    return Inches(value / 96)


def _is_markdown_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and _is_markdown_table_row(lines[index])
        and _is_markdown_table_separator(lines[index + 1])
    )


def _write_markdown_table(
    doc: Document,
    lines: list[str],
    index: int,
    style: DocxStyle,
) -> int:
    rows: list[list[str]] = []
    cursor = index
    while cursor < len(lines) and _is_markdown_table_row(lines[cursor]):
        if not _is_markdown_table_separator(lines[cursor]):
            rows.append(_split_markdown_table_row(lines[cursor]))
        cursor += 1
    if not rows:
        return cursor

    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            cell = table.rows[row_index].cells[col_index]
            value = row[col_index] if col_index < len(row) else ""
            if row_index == 0:
                _shade_cell(cell, style.table_header_bg_hex)
            _write_inline_runs(cell.paragraphs[0], value, style)
            for run in cell.paragraphs[0].runs:
                if row_index == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(*style.table_header_color_rgb)
    return cursor


def _is_markdown_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _is_markdown_table_separator(line: str) -> bool:
    cells = _split_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _split_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _ordered_list_match(line: str):
    return re.match(r"^\d+[.)]\s+(.+)$", line)


def _normalize_docx_filename(filename: str) -> str:
    name = Path(filename).name.strip() or "markdown-report.docx"
    if not name.lower().endswith(".docx"):
        name = f"{name}.docx"
    return name


def _write_list_item(doc: Document, item, style: DocxStyle, *, ordered: bool, level: int) -> None:
    paragraph_style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=paragraph_style)
    p.paragraph_format.left_indent = Cm(0.5 * level)
    run = p.add_run(item.text)
    _apply_run(run, style.body_font, style.body_size_pt, (51, 51, 51))
    for child in item.children:
        _write_list_item(doc, child, style, ordered=ordered, level=level + 1)


def _write_data_table(
    doc: Document,
    table_payload: dict[str, object],
    style: DocxStyle,
    *,
    caption: str | None = None,
) -> None:
    columns = table_payload.get("columns") if isinstance(table_payload, dict) else []
    rows = table_payload.get("rows") if isinstance(table_payload, dict) else []
    if not isinstance(columns, list) or not isinstance(rows, list):
        return
    col_keys = [c.get("key") if isinstance(c, dict) else c for c in columns]
    col_labels = [c.get("label") if isinstance(c, dict) else c for c in columns]
    if caption:
        _write_caption(doc, caption, style)
    doc_table = doc.add_table(rows=1 + len(rows), cols=max(1, len(col_labels)))
    doc_table.style = "Table Grid"
    for col_index, label in enumerate(col_labels or ["数据"]):
        cell = doc_table.rows[0].cells[col_index]
        _shade_cell(cell, style.table_header_bg_hex)
        run = cell.paragraphs[0].add_run(str(label))
        _apply_run(run, style.title_font, style.body_size_pt, style.table_header_color_rgb, bold=True)

    for row_index, row_data in enumerate(rows, start=1):
        values = [row_data.get(key) if isinstance(row_data, dict) else "" for key in col_keys]
        for cell, value in zip(doc_table.rows[row_index].cells, values, strict=False):
            run = cell.paragraphs[0].add_run(str(value if value is not None else ""))
            _apply_run(run, style.body_font, style.body_size_pt, (34, 34, 34))


def _write_chart(doc: Document, chart, style: DocxStyle, *, charts_dir: Path, caption: str | None) -> None:
    png = charts_dir / f"{chart.id}.png"
    if not png.exists():
        return
    doc.add_picture(str(png), width=Inches(6.0))
    _write_caption(doc, caption or chart.title, style)


def _write_caption(doc: Document, text: str, style: DocxStyle) -> None:
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(text)
    _apply_run(run, style.body_font, max(style.body_size_pt - 1, 8), (119, 119, 119))


def _apply_paragraph_alignment(paragraph, align: str) -> None:
    paragraph.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def _parse_rgb(value: str | None) -> tuple[int, int, int] | None:
    if not value:
        return None
    raw = value.strip().lstrip("#")
    if len(raw) != 6:
        return None
    try:
        return (int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))
    except ValueError:
        return None


def _write_kpi_table(doc: Document, kpis: list[dict[str, object]], style: DocxStyle) -> None:
    table = doc.add_table(rows=1 + len(kpis), cols=4)
    table.style = "Table Grid"

    headers = ["指标", "本期", "上期", "同比"]
    for col_index, text in enumerate(headers):
        cell = table.rows[0].cells[col_index]
        _shade_cell(cell, style.table_header_bg_hex)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(text)
        _apply_run(run, style.title_font, style.body_size_pt, style.table_header_color_rgb, bold=True)

    for row_index, kpi in enumerate(kpis, start=1):
        row = table.rows[row_index].cells
        unit = str(kpi.get("unit") or "")
        value = kpi.get("value")
        prev = kpi.get("previous_value")
        change_pct = kpi.get("change_pct")
        cells_text = [
            str(kpi.get("label", kpi.get("name", ""))),
            f"{value}{unit}" if value is not None else "—",
            f"{prev}{unit}" if prev is not None else "—",
            f"{change_pct:+.2f}%" if isinstance(change_pct, (int, float)) else "—",
        ]
        for cell, text in zip(row, cells_text, strict=True):
            run = cell.paragraphs[0].add_run(text)
            _apply_run(run, style.body_font, style.body_size_pt, (34, 34, 34))


def _write_footer(doc: Document, ctx: ReportContext, style: DocxStyle) -> None:
    p = doc.add_paragraph()
    text = (
        f"SwarmMind FlyReport · session {ctx.session_id} · "
        f"生成于 {ctx.generated_at.strftime('%Y-%m-%d %H:%M')}"
    )
    run = p.add_run(text)
    _apply_run(run, style.body_font, max(style.body_size_pt - 2, 8), (153, 153, 153))


def _apply_run(
    run,
    font_name: str,
    size_pt: int,
    rgb: tuple[int, int, int],
    *,
    bold: bool = False,
) -> None:
    _set_run_font(run, font_name, size_pt)
    run.font.color.rgb = RGBColor(*rgb)
    run.font.bold = bold


def _apply_docx_font_policy(doc: Document) -> None:
    for style in doc.styles:
        if not hasattr(style, "font"):
            continue
        style.font.name = DOCX_FONT_NAME
        run_properties = style.element.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            style.element.append(run_properties)
        run_fonts = run_properties.find(qn("w:rFonts"))
        if run_fonts is None:
            run_fonts = OxmlElement("w:rFonts")
            run_properties.append(run_fonts)
        run_fonts.set(qn("w:ascii"), DOCX_FONT_NAME)
        run_fonts.set(qn("w:hAnsi"), DOCX_FONT_NAME)
        run_fonts.set(qn("w:eastAsia"), DOCX_FONT_NAME)
        run_fonts.set(qn("w:cs"), DOCX_FONT_NAME)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, DOCX_FONT_NAME, run.font.size)
    for table in doc.tables:
        _apply_table_font_policy(table)


def _apply_table_font_policy(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, DOCX_FONT_NAME, DOCX_BODY_SIZE_PT)
            for nested_table in cell.tables:
                _apply_table_font_policy(nested_table)


def _set_run_font(run, font_name: str, size_pt: Any | None) -> None:
    run.font.name = font_name
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.append(run_fonts)
    run_fonts.set(qn("w:ascii"), font_name)
    run_fonts.set(qn("w:hAnsi"), font_name)
    run_fonts.set(qn("w:eastAsia"), font_name)
    run_fonts.set(qn("w:cs"), font_name)
    if size_pt is not None:
        run.font.size = size_pt if not isinstance(size_pt, (int, float)) else Pt(size_pt)


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _scope_label(kind: str) -> str:
    return {"overall": "总体", "department": "部门", "pilot": "飞手"}.get(kind, kind)


def _title_text(ctx: ReportContext) -> str:
    if ctx.title:
        return ctx.title
    return f"{ctx.filter.period.label} {_scope_label(ctx.filter.dimension.scope)}飞行报告"
