"""Tests for the markdown / pdf / docx renderers."""

from __future__ import annotations

from pathlib import Path

import pytest

from swarmmind.domains.fly_report.export import (
    DocxRenderer,
    MarkdownRenderer,
    PdfRenderer,
    RendererRouter,
    TemplateLoader,
)
from swarmmind.domains.fly_report.export.template_loader import PRESET_NAMES

from .conftest import make_block_context, make_context


@pytest.fixture
def loader() -> TemplateLoader:
    return TemplateLoader()


def _run_east_asia_font(run) -> str | None:
    from docx.oxml.ns import qn

    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:eastAsia"))


def _assert_docx_uses_fangsong(doc) -> None:
    paragraph_runs = [run for paragraph in doc.paragraphs for run in paragraph.runs if run.text.strip()]
    assert paragraph_runs
    assert all(run.font.name == "FangSong" for run in paragraph_runs)
    assert all(_run_east_asia_font(run) == "FangSong" for run in paragraph_runs)


def _assert_docx_tables_use_fangsong_11pt(doc) -> None:
    from docx.shared import Pt

    table_runs = [
        run
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    ]
    assert table_runs
    assert all(run.font.name == "FangSong" for run in table_runs)
    assert all(_run_east_asia_font(run) == "FangSong" for run in table_runs)
    assert all(run.font.size == Pt(11) for run in table_runs)


# ----------------------------------------------------------------- markdown


def test_markdown_default_renders_kpi_table(tmp_path: Path, loader):
    renderer = MarkdownRenderer(template_loader=loader)
    artifact = renderer.render(make_context(), output_dir=tmp_path)

    assert artifact.output_format == "markdown"
    assert artifact.template_ref == "default"
    assert artifact.warnings == []

    content = Path(artifact.artifact_path).read_text(encoding="utf-8")
    assert "飞行报告" in content
    assert "飞行任务数" in content
    assert "+25.49%" in content
    # KPI with no previous value should render an em dash, not crash.
    assert "算法告警" in content


@pytest.mark.parametrize("preset", PRESET_NAMES)
def test_markdown_each_preset_renders(tmp_path: Path, loader, preset):
    renderer = MarkdownRenderer(template_loader=loader)
    artifact = renderer.render(
        make_context(),
        output_dir=tmp_path,
        template_ref=f"preset:{preset}",
    )
    assert artifact.template_ref == f"preset:{preset}"
    content = Path(artifact.artifact_path).read_text(encoding="utf-8")
    assert "飞行" in content


def test_markdown_renders_report_blocks_in_order(tmp_path: Path, loader):
    renderer = MarkdownRenderer(template_loader=loader)
    artifact = renderer.render(make_block_context(), output_dir=tmp_path)

    content = Path(artifact.artifact_path).read_text(encoding="utf-8")
    assert "# 自定义飞行报告" in content
    assert "本周飞行任务整体稳定" in content
    assert "任务趋势" in content
    assert "![飞行趋势](charts/flight-trend.png)" in content
    assert "部门明细" in content
    assert "复盘夜间航线" in content
    assert content.index("本周飞行任务整体稳定") < content.index("任务趋势")
    assert (tmp_path / "charts" / "flight-trend.png").exists()


# ----------------------------------------------------------------- pdf


def test_pdf_falls_back_to_html_when_weasyprint_missing(
    tmp_path: Path, loader, monkeypatch
):
    # Force the optional import to fail even if weasyprint happens to be
    # installed, so the test asserts the documented fallback.
    import builtins

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "weasyprint":
            raise ImportError("simulated absence")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    renderer = PdfRenderer(template_loader=loader)
    artifact = renderer.render(make_context(), output_dir=tmp_path)

    assert artifact.output_format == "pdf"
    assert artifact.artifact_path.endswith(".html")
    assert any("WeasyPrint" in w for w in artifact.warnings)
    html = Path(artifact.artifact_path).read_text(encoding="utf-8")
    assert "<h1>" in html
    assert "飞行任务数" in html


@pytest.mark.parametrize("preset", PRESET_NAMES)
def test_pdf_each_preset_renders_html(tmp_path: Path, loader, monkeypatch, preset):
    import builtins

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "weasyprint":
            raise ImportError("simulated absence")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    renderer = PdfRenderer(template_loader=loader)
    artifact = renderer.render(
        make_context(),
        output_dir=tmp_path,
        template_ref=f"preset:{preset}",
    )
    assert artifact.template_ref == f"preset:{preset}"
    html = Path(artifact.artifact_path).read_text(encoding="utf-8")
    assert "飞行" in html


def test_pdf_html_fallback_renders_report_blocks(tmp_path: Path, loader, monkeypatch):
    import builtins

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "weasyprint":
            raise ImportError("simulated absence")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    renderer = PdfRenderer(template_loader=loader)
    artifact = renderer.render(make_block_context(), output_dir=tmp_path)

    html = Path(artifact.artifact_path).read_text(encoding="utf-8")
    assert "自定义飞行报告" in html
    assert "本周飞行任务整体稳定" in html
    assert "charts/flight-trend.png" in html
    assert "部门明细" in html
    assert (tmp_path / "charts" / "flight-trend.png").exists()


# ----------------------------------------------------------------- docx


def test_docx_default_writes_valid_file(tmp_path: Path, loader):
    from docx import Document

    renderer = DocxRenderer(template_loader=loader)
    artifact = renderer.render(make_context(), output_dir=tmp_path)

    assert artifact.output_format == "docx"
    assert artifact.template_ref == "default"
    out = Path(artifact.artifact_path)
    assert out.exists() and out.stat().st_size > 0

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "飞行报告" in text
    _assert_docx_uses_fangsong(doc)
    _assert_docx_tables_use_fangsong_11pt(doc)
    # KPI table should have rendered.
    assert any(
        "飞行任务数" in cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )


@pytest.mark.parametrize("preset", PRESET_NAMES)
def test_docx_each_preset_writes_valid_file(tmp_path: Path, loader, preset):
    from docx import Document

    renderer = DocxRenderer(template_loader=loader)
    artifact = renderer.render(
        make_context(),
        output_dir=tmp_path,
        template_ref=f"preset:{preset}",
    )
    assert artifact.template_ref == f"preset:{preset}"
    doc = Document(artifact.artifact_path)
    assert any("飞行报告" in p.text for p in doc.paragraphs)


def test_docx_renders_report_blocks(tmp_path: Path, loader):
    from docx import Document

    renderer = DocxRenderer(template_loader=loader)
    artifact = renderer.render(make_block_context(), output_dir=tmp_path)

    doc = Document(artifact.artifact_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "自定义飞行报告" in text
    assert "本周飞行任务整体稳定" in text
    assert "任务趋势" in text
    assert "复盘夜间航线" in text
    assert "飞行任务数" in table_text
    assert "一中队" in table_text
    _assert_docx_uses_fangsong(doc)
    _assert_docx_tables_use_fangsong_11pt(doc)
    assert (tmp_path / "charts" / "flight-trend.png").exists()


# ----------------------------------------------------------------- router


def test_router_dispatches_by_output_format(tmp_path: Path):
    router = RendererRouter()
    md = router.render(
        make_context(), output_format="markdown", output_dir=tmp_path
    )
    assert md.artifact_path.endswith(".md")

    docx = router.render(
        make_context(),
        output_format="docx",
        output_dir=tmp_path,
        template_ref="preset:minimal",
    )
    assert docx.artifact_path.endswith(".docx")
    assert docx.template_ref == "preset:minimal"


def test_router_renders_markdown_directly_to_docx(tmp_path: Path):
    import matplotlib
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    image_path = tmp_path / "chart.png"
    plt.figure(figsize=(2, 1))
    plt.plot([1, 2, 3], [2, 1, 3])
    plt.savefig(image_path)
    plt.close()

    router = RendererRouter()
    markdown = f"""# 飞行周报

::: {{align=center}}
本周 **飞行任务** 整体稳定。
:::

::: {{align=right}}
[右对齐提醒]{{color=#d62728}}
:::

- 完成巡查 20 次
- 异常任务 2 次

| 指标 | 数值 |
| --- | --- |
| 飞行次数 | 20 |
| 告警数量 | 8 |

::: {{align=right}}
| 右对齐表格 | 数值 |
| --- | ---: |
| 处置率 | 91.4% |
:::

::: {{.table-pair}}
::: {{align=left}}
| 左表 | 值 |
| --- | ---: |
| A | 1 |
:::

::: {{align=right}}
| 右表 | 值 |
| --- | ---: |
| B | 2 |
:::
:::

![趋势图]({image_path}){{width=5cm align=center}}
"""
    artifact = router.render_markdown_to_docx(
        markdown,
        output_dir=tmp_path,
        filename="weekly-report",
        title="武义飞行服务平台",
    )

    assert artifact.output_format == "docx"
    assert artifact.artifact_path.endswith("weekly-report.docx")
    assert artifact.template_ref == "markdown:default"

    doc = Document(artifact.artifact_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "武义飞行服务平台" in text
    assert "飞行周报" in text
    assert "完成巡查 20 次" in text
    assert "飞行次数" in table_text
    assert "20" in table_text
    _assert_docx_uses_fangsong(doc)
    _assert_docx_tables_use_fangsong_11pt(doc)
    assert any(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in doc.paragraphs)
    assert any(p.alignment == WD_ALIGN_PARAGRAPH.RIGHT for p in doc.paragraphs)
    assert len(doc.inline_shapes) == 1
    drawing_paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//w:drawing")]
    assert drawing_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert any(table.alignment == WD_TABLE_ALIGNMENT.RIGHT for table in doc.tables)
    paired = [table for table in doc.tables if len(table.rows) == 1 and len(table.columns) == 2]
    assert paired
    assert len(paired[0].rows[0].cells[0].tables) == 1
    assert len(paired[0].rows[0].cells[1].tables) == 1


def test_router_rejects_unknown_format(tmp_path: Path):
    router = RendererRouter()
    with pytest.raises(ValueError, match="Unsupported output_format"):
        router.render(
            make_context(),
            output_format="xls",  # type: ignore[arg-type]
            output_dir=tmp_path,
        )
